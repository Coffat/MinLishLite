#!/usr/bin/env python3
"""Fill docs/BaoCaoKetQuaDoAn.docx with MinLishLite project report content."""

from pathlib import Path

from docx import Document
from docx.shared import Inches

ROOT = Path(__file__).resolve().parent
DOC_PATH = ROOT / "BaoCaoKetQuaDoAn.docx"

SECTION_1 = """MinLish Lite là ứng dụng Android hỗ trợ học từ vựng tiếng Anh bằng flashcard và hệ thống lặp lại ngắt quãng (Spaced Repetition System – SRS). Ứng dụng theo hướng offline-first: dữ liệu chính lưu cục bộ bằng Room Database; mạng chỉ dùng để tra từ điển và dịch thuật.

Các chức năng chính đã hoàn thành:

• Onboarding & hồ sơ (tính năng cơ bản): Tiếp tục với tư cách Guest hoặc đăng nhập Google (Play Services Auth), thiết lập tên/mục tiêu/trình độ, lưu UserEntity cục bộ.

• Quản lý bộ từ: Tạo, sửa, xóa bộ từ; gắn tag; tìm kiếm và lọc theo tag (DeckListScreen, AddEditDeckScreen).

• Quản lý từ vựng: Thêm, sửa, xóa từ; xem chi tiết; lọc theo Tất cả / Đến hạn / Đã học (WordDetailScreen, AddEditWordScreen).

• Tra từ điển API: Tự động lấy phiên âm UK/US, audio, định nghĩa qua dictionaryapi.dev; dịch sang tiếng Việt qua MyMemory API (DictionaryRepository).

• Import/Export CSV: Nhập và xuất từ vựng theo bộ từ qua CsvHelper.

• Học flashcard: Lật thẻ, tự đánh giá Again / Hard / Good / Easy (StudyScreen, Flashcard).

• SRS (Spaced Repetition System): SrsCalculator tính lịch ôn dựa trên hệ số độ dễ (easeFactor). Khoảng cách ôn = số ngày cơ bản × easeFactor (Again: ngay; Hard: 1 × easeFactor ngày; Good: 3 × easeFactor ngày; Easy: 7 × easeFactor ngày). Ôn tốt liên tiếp → easeFactor tăng → khoảng cách ngày càng dài, đúng với đường cong lãng quên.

• Ôn tập hôm nay: Danh sách từ có nextReviewAt <= thời điểm hiện tại (ReviewTodayScreen).

• Theo dõi tiến độ: Streak, độ chính xác, retention, biểu đồ hoạt động 7 ngày, thành tựu (ProgressScreen).

• Cài đặt & nhắc học: Cấu hình số từ mới/ngày, bật/tắt nhắc nhở, giờ nhắc qua DataStore; banner nhắc trên Home; notification nhắc học thật sự qua WorkManager (StudyReminderWorker chạy đúng giờ mỗi ngày) (SettingsScreen).

• Phát âm: Phát audio URL từ API qua MediaPlayer (PronunciationAudioPlayer) trên flashcard và màn thêm từ."""

SECTION_3 = """Kiến trúc phần mềm:
• MVVM (Model-View-ViewModel), Single Activity (MainActivity).
• Luồng dữ liệu: Composable UI → ViewModel → Repository → Room DAO / Retrofit API.
• StateFlow quản lý trạng thái UI theo hướng Unidirectional Data Flow.

Design Pattern:
• Repository Pattern — mọi truy cập dữ liệu qua repository, không gọi DAO/API từ Composable.
• Manual Dependency Injection qua AppContainer.
• Sealed class StudyMode cho các chế độ học.

Database:
• Room Database phiên bản 4 (minlish_database): users, decks, words, review_history.
• Quan hệ: decks 1-N words, words 1-N review_history (CASCADE on delete).
• Migration 1→4; index trên deckId và nextReviewAt.

Unit Test:
• 9 file unit test: SrsCalculatorTest (bao gồm test kiểm tra easeFactor ảnh hưởng interval), ProgressCalculatorTest, WordValidatorTest, StudyViewModelTest, ReviewTodayViewModelTest, AddEditWordViewModelTest, StudyRepositoryTest, WordRepositoryTest, DictionaryRepositoryTest.
• 4 file instrumented test cho DAO: WordDaoTest, DeckDaoTest, UserDaoTest, ReviewHistoryDaoTest.

Performance Optimization:
• Kotlin Coroutines và Flow cho thao tác bất đồng bộ.
• Room reactive queries; LazyColumn cho danh sách dài; composable nhỏ, tách biệt.

Security:
• Không hardcode API key; dữ liệu người dùng lưu cục bộ.
• Quyền INTERNET, POST_NOTIFICATIONS, RECEIVE_BOOT_COMPLETED.

Thư viện / Framework:
• Kotlin 2.2, Jetpack Compose Material 3, Navigation Compose.
• Room (KSP), Retrofit + OkHttp + Gson, DataStore Preferences.
• Google Play Services Auth.
• WorkManager (work-runtime-ktx): StudyReminderWorker (CoroutineWorker) + StudyReminderScheduler lên lịch nhắc học định kỳ mỗi 24 giờ."""

SECTION_4 = """Các chức năng ngoài yêu cầu đề bài cơ bản:

1. Tích hợp Dictionary API + dịch tự động — tự điền phiên âm UK/US, nghĩa, ví dụ qua dictionaryapi.dev và MyMemory API khi thêm từ.

2. Phát âm từ audio API — phát file audio UK/US trên flashcard và màn thêm từ thông qua PronunciationAudioPlayer (MediaPlayer).

3. Hệ thống achievements — 5 huy hiệu trong ProgressCalculator (lần ôn đầu, streak 7 ngày, 100 từ, accuracy 80%, retention 70%).

4. Import/Export CSV — nhập/xuất bộ từ nhanh qua CsvHelper, hỗ trợ chia sẻ dữ liệu."""

SECTION_6 = """Khó khăn gặp phải:

• Đồng bộ nhiều Flow trong ViewModel (combine) khi cần gộp dữ liệu từ nhiều nguồn — đòi hỏi hiểu reactive programming.
• Migration Room khi mở rộng WordEntity (từ v1 lên v4) — cần viết migration script cẩn thận để không mất dữ liệu.
• API công khai (dictionary, dịch) bị rate limit và lỗi mạng — phải xử lý Result và hiển thị trạng thái lỗi trên UI.
• Phân chia trách nhiệm 2 thành viên theo feature nhưng vẫn phải nắm phần chung (navigation, DI, database).
• Triển khai thuật toán SRS kết hợp easeFactor — phải hiểu rõ SM-2 để tính khoảng cách ôn tăng dần đúng cách.
• WorkManager: tính initialDelay chính xác đến giờ nhắc tiếp theo và xử lý trường hợp giờ đã qua trong ngày.

Bài học kinh nghiệm:

• MVVM kết hợp Repository giúp tách logic khỏi UI và dễ viết unit test.
• Đặt thuật toán SRS trong SrsCalculator thuần (không phụ thuộc Android) giúp kiểm thử độc lập và dễ mở rộng.
• CoroutineWorker cho phép code background chạy trong coroutine, tránh ANR và dễ đọc hơn AsyncTask.
• Hướng offline-first phù hợp đồ án Android cơ bản, app vẫn dùng được khi mất mạng.
• State hoisting và StateFlow giúp UI Compose dự đoán được và dễ bảo trì."""

SECTION_7 = """Đề xuất hướng phát triển:

1. Chuyển sang Hilt hoặc Koin — thay manual DI trong AppContainer bằng framework Dependency Injection chuẩn.

2. Đồng bộ cloud — backup/restore bộ từ trên nhiều thiết bị (Firebase Firestore hoặc backend riêng).

3. TTS fallback — phát âm bằng Text-to-Speech khi API không có URL audio.

4. Thống kê nâng cao — heatmap hoạt động học theo năm, export báo cáo học tập dạng PDF.

5. Dark mode đầy đủ trên toàn ứng dụng.

6. Gamification mở rộng — bảng xếp hạng, chia sẻ thành tựu lên mạng xã hội."""

COMPLETION_TABLE = [
    ("Tạo bộ từ vựng", "100%"),
    ("Thêm từ vựng", "100%"),
    ("Học/ôn từ vựng", "100%"),
    ("Quản lý người dùng", "95%"),
    ("Import/Export", "100%"),
    ("Ôn tập thông minh", "100%"),
    ("Theo dõi tiến độ", "100%"),
    ("Thông báo nhắc học", "100%"),
    ("Chức năng khác", "100%"),
]

RUBRIC_TABLE = [
    ("Bài nộp", "5", "5", "Slide MinLishLite_Presentation.pptx, video demo, source code GitHub"),
    ("Chất lượng sản phẩm", "30", "28", "MVVM rõ, UI ổn, 13 test, WorkManager Worker hoàn chỉnh; trừ điểm nhỏ DI thủ công"),
    ("Chức năng cơ bản", "30", "30", "Deck/Word/Study/Onboarding đầy đủ"),
    ("Chức năng nâng cao", "25", "25", "SRS dùng easeFactor thật, WorkManager notification đúng giờ, CSV, progress analytics"),
    ("Sáng tạo & mở rộng", "10", "8", "Dictionary API + phát âm audio + achievements; chưa có cloud sync"),
    ("TỔNG ĐIỂM", "100", "96", "SRS và notification đã hoàn thiện, nâng từ 92 lên 96"),
]


def set_paragraph_text(paragraph, text: str) -> None:
    for run in paragraph.runs:
        run.text = ""
    if paragraph.runs:
        paragraph.runs[0].text = text
    else:
        paragraph.add_run(text)


def delete_paragraph(paragraph) -> None:
    element = paragraph._element
    element.getparent().remove(element)


def find_paragraph(doc: Document, prefix: str):
    for paragraph in doc.paragraphs:
        if paragraph.text.strip().startswith(prefix):
            return paragraph
    return None


def add_figure_before(ref_paragraph, caption: str, image_path: Path, width_inches: float = 5.0) -> None:
    if not image_path.exists():
        return
    cap = ref_paragraph.insert_paragraph_before(caption)
    cap.alignment = 1
    img_p = ref_paragraph.insert_paragraph_before("")
    img_p.alignment = 1
    run = img_p.add_run()
    run.add_picture(str(image_path), width=Inches(width_inches))


def fill_header_tables(doc: Document) -> None:
    header = doc.tables[0]
    header.rows[0].cells[0].text = "Môn học: Lập trình di động\nLớp:"
    header.rows[0].cells[1].text = "Mã nhóm: 21\nSố lượng thành viên: 2"

    members = doc.tables[1]
    members.rows[1].cells[0].text = "23110329"
    members.rows[1].cells[1].text = "Vũ Toàn Thắng (nhóm trưởng)"
    members.rows[1].cells[2].text = "22110334"
    members.rows[1].cells[3].text = "Nguyễn Tuấn Huy"
    members.rows[1].cells[4].text = "-"
    members.rows[1].cells[5].text = "-"
    members.rows[3].cells[0].text = "-"
    members.rows[3].cells[1].text = "-"
    members.rows[3].cells[2].text = "-"
    members.rows[3].cells[3].text = "-"


def fill_completion_table(doc: Document) -> None:
    table = doc.tables[2]
    for row_idx, (feature, pct) in enumerate(COMPLETION_TABLE, start=1):
        table.rows[row_idx].cells[1].text = pct


def fill_rubric_table(doc: Document) -> None:
    table = doc.tables[3]
    for row_idx, row_data in enumerate(RUBRIC_TABLE, start=1):
        for col_idx, value in enumerate(row_data):
            table.rows[row_idx].cells[col_idx].text = value


def remove_existing_diagrams(doc: Document) -> None:
  to_delete = []
  for paragraph in doc.paragraphs:
    text = paragraph.text.strip()
    if text.startswith("Hình 1:") or text.startswith("Hình 2:"):
      to_delete.append(paragraph)
      continue
    if not text and paragraph.runs:
      for run in paragraph.runs:
        if run._element.xpath(".//a:blip"):
          to_delete.append(paragraph)
          break
  for paragraph in to_delete:
    delete_paragraph(paragraph)


def fill_sections(doc: Document) -> None:
  section_map = {
    "Các chức năng chính đã hoàn thành": SECTION_1,
    "MinLish Lite là ứng dụng Android": SECTION_1,
    "- Kiến trúc phần mềm": SECTION_3,
    "Kiến trúc phần mềm:": SECTION_3,
    "Liệt kê các chức năng ngoài yêu cầu": SECTION_4,
    "Các chức năng ngoài yêu cầu đề bài": SECTION_4,
    "Mô tả các khó khăn gặp phải": SECTION_6,
    "Khó khăn gặp phải:": SECTION_6,
    "Đề xuất các tính năng hoặc cải tiến": SECTION_7,
    "Đề xuất hướng phát triển:": SECTION_7,
  }
  for paragraph in doc.paragraphs:
    text = paragraph.text.strip()
    for prefix, content in section_map.items():
      if text.startswith(prefix):
        set_paragraph_text(paragraph, content)
        break


def insert_diagrams(doc: Document) -> None:
  section_3_para = find_paragraph(doc, "Kiến trúc phần mềm:")
  if section_3_para is None:
    return
  remove_existing_diagrams(doc)
  navflow = ROOT / "navflow.png"
  datacore = ROOT / "datacore.png"
  add_figure_before(section_3_para, "Hình 1: Luồng điều hướng ứng dụng", navflow)
  add_figure_before(section_3_para, "Hình 2: Mô hình dữ liệu (Room Database)", datacore)


def main() -> None:
    doc = Document(str(DOC_PATH))
    fill_header_tables(doc)
    fill_sections(doc)
    fill_completion_table(doc)
    fill_rubric_table(doc)
    insert_diagrams(doc)
    doc.save(str(DOC_PATH))
    print(f"Saved: {DOC_PATH}")


if __name__ == "__main__":
    main()
